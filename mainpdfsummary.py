import streamlit as st
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
from langchain_community.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.summarize import load_summarize_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
import traceback
from docx import Document
from langchain.docstore.document import Document


# Load API key from Streamlit secrets
google_api_key = st.secrets.get("GOOGLE_API_KEY")
if not google_api_key:
    st.error("GOOGLE_API_KEY not found in Streamlit secrets.")
    st.stop()
genai.configure(api_key=google_api_key)

def get_document_text(docs):
    text = ""
    for doc in docs:
        try:
            if doc.name.lower().endswith(".pdf"):
                pdf_reader = PdfReader(doc)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text
            elif doc.name.lower().endswith(".txt"):
                text += doc.getvalue().decode("utf-8")
            elif doc.name.lower().endswith(".docx"):
                document = Document(doc)
                for paragraph in document.paragraphs:
                    text += paragraph.text + "\n"
            else:
                st.warning(f"Unsupported file type: {doc.name}. Only PDF, TXT and DOCX are supported.")
                continue
        except Exception as e:
            st.error(f"Error reading document {doc.name}: {str(e)}")
            continue
    return text

def get_text_chunks(text):
    if not text.strip():
        return []
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=10000,
        chunk_overlap=1000
    )
    chunks = text_splitter.split_text(text)
    return chunks

def generate_summary(text_chunks, temperature):
    if not text_chunks:
        st.error("No text chunks provided for summarization.")
        return ""

    try:
        model = ChatGoogleGenerativeAI(
            model="gemini-1.5-pro-latest",
            temperature=temperature
        )

        prompt_template = """
        Write a concise summary of the following text:
        {text}

        SUMMARY:
        """
        prompt = PromptTemplate(template=prompt_template, input_variables=["text"])

        chain = load_summarize_chain(model, chain_type="stuff", prompt=prompt)

        docs = [Document(page_content=chunk) for chunk in text_chunks]

        summary = chain.invoke({"input_documents": docs})
        return summary["output_text"]

    except Exception as e:
        st.error(f"Failed to generate summary: {str(e)}")
        return ""

def delete_summary(doc_name):
    summary_dir = f"summary_{doc_name}"
    pdf_path = os.path.join(summary_dir, f"{doc_name}_summary.pdf")

    if os.path.exists(pdf_path):
        os.remove(pdf_path)
        st.success(f"Summary deleted from {pdf_path}")
    else:
        st.warning("Summary file not found.")

def create_vector_store(text_chunks):
    if not text_chunks:
        raise ValueError("No text chunks provided for vector store creation")
    embeddings = GoogleGenerativeAIEmbeddings(model="models/text-embedding-004")
    vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
    vector_store.save_local("faiss_index")
    st.success("Vector store created successfully!")

def main():
    #st.set_page_config("DocuMind AI", page_icon="📄", layout="wide")
    
    # Custom CSS for better styling
    st.markdown("""
    <style>
    .main {
        max-width: 1200px;
    }
    .sidebar .sidebar-content {
        background-color: #f8f9fa;
    }
    .stButton>button {
        background-color: #4CAF50;
        color: white;
        border-radius: 5px;
        padding: 10px 24px;
    }
    .stButton>button:hover {
        background-color: #45a049;
    }
    .file-uploader {
        padding: 20px;
        border-radius: 10px;
        background-color: #f0f2f6;
    }
    .instructions {
        padding: 15px;
        background-color: #e9f7ef;
        border-radius: 10px;
        margin-bottom: 20px;
    }
    </style>
    """, unsafe_allow_html=True)

    # Main layout columns
    col1, col2 = st.columns([3, 1])  # Wider main content, narrower sidebar

    with col1:
        st.header("📄 DocuMind AI - Document Summarizer")
        st.markdown("""
        **Transform lengthy documents into concise summaries with AI-powered analysis.**
        Upload your files and get instant summaries tailored to your needs.
        """)

        if st.session_state.get("summary"):
            st.subheader("Your Document Summary")
            with st.expander("View Summary", expanded=True):
                st.write(st.session_state.summary)
            
            if st.button("Clear Summary"):
                st.session_state.summary = None
                st.session_state.doc_name = None
                st.rerun()
        else:
            st.info("Upload your documents to generate a summary")

    with col2:
        st.markdown('<div class="instructions">', unsafe_allow_html=True)
        st.subheader("How to Use")
        st.markdown("""
        1. Upload documents (PDF, DOCX, TXT)
        2. Adjust creativity level
        3. Click "Generate Summary"
        4. View and download your summary
        """)
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="file-uploader">', unsafe_allow_html=True)
        document_docs = st.file_uploader(
            "Upload Documents",
            accept_multiple_files=True,
            type=["pdf", "txt", "docx"],
            help="Upload one or more documents to summarize"
        )

        temperature = st.slider(
            "Creativity Level",
            min_value=0.0,
            max_value=1.0,
            value=0.3,
            step=0.01,
            help="Higher values produce more creative summaries"
        )

        if st.button("Generate Summary", key="generate_btn"):
            if not document_docs:
                st.warning("Please upload at least one document file.")
            else:
                with st.spinner("Analyzing documents..."):
                    try:
                        raw_text = get_document_text(document_docs)

                        if not raw_text.strip():
                            st.error("No text could be extracted from the document(s).")
                            return

                        text_chunks = get_text_chunks(raw_text)

                        if not text_chunks:
                            st.error("Failed to split the text into meaningful chunks.")
                            return

                        summary = generate_summary(text_chunks, temperature)

                        if summary:
                            st.session_state.summary = summary
                            st.session_state.doc_name = document_docs[0].name.split(".")[0]
                            create_vector_store(text_chunks)
                            st.rerun()

                    except Exception as e:
                        st.error(f"Error: {str(e)}")
        st.markdown('</div>', unsafe_allow_html=True)

if __name__ == "__main__":
    if "summary" not in st.session_state:
        st.session_state.summary = None
    if "doc_name" not in st.session_state:
        st.session_state.doc_name = None
    main()
def app():
    main()
